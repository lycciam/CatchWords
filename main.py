import requests
from bs4 import BeautifulSoup
import json
from docx import Document

workfile = 'F:\\work_of_spyder\\'
document = Document()

main_url = 'http://tieba.baidu.com/p/3618228828'

header = {
    'User-Agent':
        'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/48.0.2564.109 Safari/537.36'}

'''
    判断一个贴吧帖子的页面数量，如果大于一，则继续玩，不然就不玩。
'''

def get_content_of_floor(soup_floor, Docx):
    """
    获得每个楼层的作者，发帖时间，以及内容
    :param soup_floor: 楼层的soup节点
    :param Docx: document的句柄
    """
    information = soup_floor.get('data-field', 'N/A')
    if '&quot;' in information:
        information = information.replace('&quot;', '"')
    information = json.loads(information)
    time = str(information['content']['date'])
    author = information['author']['user_name']
    content = soup_floor.select('div.p_content.p_content_nameplate > cc > div')[0].get_text()
    Docx.add_paragraph('date:' + time)
    Docx.add_paragraph('author:' + author)
    Docx.add_paragraph('content:' + content)
    Docx.add_paragraph('---------楼层间隔符----------')


page = requests.get(main_url, header)
soup = BeautifulSoup(page.text, 'lxml')
page_num = soup.select('#thread_theme_5 > div.l_thread_info > ul > li:nth-of-type(2) > '
                       'span:nth-of-type(2)')[0].get_text()
init = soup.select('div.p_postlist > div')[0]

# 获取基本信息
time = init.get('data-field', 'N/A')
time = json.loads(time)
title = soup.select('title')[0].get_text()
document.add_paragraph('url:' + main_url)
document.add_paragraph("title:" + title)
document.add_paragraph("time:" + str(time['content']['date']))
document.add_paragraph("author:" + time['author']['user_name'])
document.add_paragraph('post_id:' + str(time['content']['post_id']))
document.add_paragraph('回复数量:' + soup.select('#thread_theme_5 > div.l_thread_info > ul > li:nth-of-type(2) > '
                                             'span:nth-''of-type(1)')[0].get_text())
document.add_paragraph('总页数:' + soup.select('#thread_theme_5 > div.l_thread_info > ul > li:nth-of-type(2) > span:nth-'
                                            'of-type(2)')[0].get_text())
# 好像少了个一楼的内容？
document.add_paragraph('---------楼层间隔符----------')

urls = []
if int(page_num) == 1:
    urls.append(main_url)
else:
    for i in range(int(page_num)):
        url = main_url + '?pn=' + str(i + 1)
        urls.append(url)

for each_url in urls:
    page_each = requests.get(each_url, header)
    soup_each = BeautifulSoup(page_each.text, 'lxml')
    print(len(soup_each.select('div.p_postlist > div.l_post.j_l_post.l_post_bright')))
    soups = soup_each.select('div.p_postlist > div.l_post.j_l_post.l_post_bright')
    for each_soup in soups:
        get_content_of_floor(each_soup, document)

# save file
document.save(workfile + title + '.docx')



