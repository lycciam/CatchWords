import requests
from bs4 import BeautifulSoup
import json
from docx import Document
import logging.handlers

# 声明日志区
# 日志的区域需要另外去写
log_file = 'F:\\work_of_spyder\\Logfile\\react.log'
handler = logging.handlers.RotatingFileHandler(log_file, maxBytes=10 * 1024 * 1024, backupCount=5)
fmt = '%(asctime)s - %(filename)s:%(lineno)s - %(levelname)s - %(message)s - %(msecs)s'
formatter = logging.Formatter(fmt)
handler.setFormatter(formatter)
logger = logging.getLogger('done')
logger.addHandler(handler)
logger.setLevel(logging.DEBUG)

workfile = 'F:\\work_of_spyder\\'
document = Document()

header = {
    'User-Agent':
        'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/48.0.2564.109 Safari/537.36'}

def rpc_char1(text):
    """
        将目标文本替换true
    """
    text = text.replace('true', "1")
    text = text.replace('True', "1")
    return text


def get_content_of_floor(soup_floor, docx):
    """
    获得每个楼层的作者，发帖时间，以及内容
    :param soup_floor: 楼层的soup节点
    :param docx: document的句柄
    """
    information = soup_floor.get('data-field', 'N/A')
    if '&quot;' in information:
        information = information.replace('&quot;', '"')
    information = rpc_char1(information)
    information = json.loads(information)
    time = str(information['content'].get('date', 'N/A'))
    author = information['author'].get('user_name', 'N/A')
    content = soup_floor.select('div > div > cc > div')[0].get_text()
    docx.add_paragraph('date:' + time)
    docx.add_paragraph('author:' + author)
    docx.add_paragraph('content:' + content)
    docx.add_paragraph('---------楼层间隔符----------')


def rpc_char(text):
    """
        将目标文本替换掉不能作为文件名的字符
    """
    char_set = ["/", "\\", "?", "*", "<", ">", "|", '"', "！"]
    for i in char_set:
        text = text.replace(i, " ")
    return text




def isnone_tieba(url):
    """
    判断此贴吧的url是不是空的
    :param url: url呗。不过以后应该是要写个判定格式的。
    """
    page = requests.get(url, header)
    soup = BeautifulSoup(page.text, 'lxml')
    if soup.select('title')[0].get_text() == '贴吧404':
        logger.info('无效url:' + url)
    elif soup.select('title')[0].get_text() == '百度贴吧':
        logger.info('无效url:' + url)
    else:
        test = 1
        return test


def get_tiezi_info(url, headers, work_file):
    """
    主要的工作区
    :param url: 目标url
    :param headers:     浏览器参数
    :param work_file:   保存的工作文档路径
    """
    page = requests.get(url, headers)
    soup = BeautifulSoup(page.text, 'lxml')
    page_num = soup.select('#thread_theme_5 > div.l_thread_info > ul > li:nth-of-type(2) > '
                           'span:nth-of-type(2)')[0].get_text()
    init = soup.select('div.p_postlist > div')[0]

    # 获取基本信息
    time = init.get('data-field', 'N/A')
    time = json.loads(time)
    title = soup.select('title')[0].get_text()
    document.add_paragraph('url:' + url)
    document.add_paragraph("title:" + title)
    document.add_paragraph("time:" + str(time['content'].get('date', 'N/A')))
    document.add_paragraph("author:" + time['author'].get('user_name', 'N/A'))
    document.add_paragraph('post_id:' + str(time['content'].get('post_id', 'N/A')))
    document.add_paragraph('回复数量:' + soup.select('#thread_theme_5 > div.l_thread_info > ul > li:nth-of-type(2) > '
                                                 'span:nth-''of-type(1)')[0].get_text())
    document.add_paragraph(
        '总页数:' + soup.select('#thread_theme_5 > div.l_thread_info > ul > li:nth-of-type(2) > span:nth-'
                             'of-type(2)')[0].get_text())
    if len(soup.select('a.j_plat_picbox.plat_picbox')) != 0:
        document.add_paragraph('贴吧名:' + soup.select('a.j_plat_picbox.plat_picbox')[0].get_text())
    else:
        document.add_paragraph('贴吧名:' + soup.select('a.card_title_fname')[0].get_text())
    document.add_paragraph('---------楼层间隔符----------')

    urls = []
    if int(page_num) == 1:
        urls.append(url)
    else:
        for i in range(int(page_num)):
            url = url + '?pn=' + str(i + 1)
            urls.append(url)

    for each_url in urls:
        page_each = requests.get(each_url, header)
        soup_each = BeautifulSoup(page_each.text, 'lxml')
        soups = soup_each.select('div.p_postlist > div.l_post.j_l_post.l_post_bright')
        for each_soup in soups:
            get_content_of_floor(each_soup, document)

    # save file

    title = rpc_char(title)

    document.save(work_file + title + '.docx')
