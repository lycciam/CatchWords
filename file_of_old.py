import requests
from bs4 import BeautifulSoup
import re
from docx import Document

url = "http://tieba.baidu.com/p/811767739"

header = {
    'User-Agent':
        'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/48.0.2564.109 Safari/537.36'}
# 首次载入
page = requests.get(url, header)
soup = BeautifulSoup(page.text, 'lxml')

# 初始化文档
document = Document()

# 获得帖子的基本信息
print('url:', url)
print("title:", soup.title.contents[0])
author = soup.select("#j_p_postlist > div.l_post.j_l_post.l_post_bright.noborder > div.d_author > ul > li.d_name > a")
print("author:", author[0].text)
num_page = soup.select('div.thread_theme_5 > div > ul[class~=l_posts_num] > li[class~=l_reply_num] > span[class~=red]')
print(num_page[1].text)

document.add_heading('url:' + url, 6)
document.add_paragraph("title:" + soup.title.contents[0])
document.add_paragraph("pages:" + num_page[1].text)

# 获取一楼的信息；
# 作者，发布时间，子贴id，回复数量，帖子内容
document.add_paragraph("")
document.add_paragraph("author:" + author[0].text)

Infmn = soup.find_all('div', class_='l_post j_l_post l_post_bright noborder ')[0].get("data-field")
date = re.compile(r'\"date\"\:\"(.*?)\"\,')
floor = re.compile(r'\"post_no"\:(.*?)\,')
post_id = re.compile(r'\"post_id\"\:(.*?)\,')
comment_num = re.compile(r'\"comment_num\"\:(.*?)\,')
document.add_paragraph('date:' + date.findall(Infmn)[0])
document.add_paragraph('floor:' + floor.findall(Infmn)[0])
document.add_paragraph('post_id:' + post_id.findall(Infmn)[0])
document.add_paragraph('comment_num:' + comment_num.findall(Infmn)[0])
document.add_paragraph('contents:')
neirong = (soup.find_all('div', class_='l_post j_l_post l_post_bright noborder ')[0]). \
    find('div', class_='d_post_content j_d_post_content  clearfix')
for child in neirong:
    document.add_paragraph(str(child))
document.add_paragraph('---------楼层间隔符----------')

# 获得第一页内除了第一帖子之外的所有帖子的数量
num_tiezi = soup.find_all('div', class_='l_post j_l_post l_post_bright ')

if len(num_tiezi) == 0:
    pass
else:
    for i in range(len(num_tiezi)):
        Infmn = soup.find_all('div', class_='l_post j_l_post l_post_bright ')[i].get('data-field')
        date = re.compile(r'\"date\"\:\"(.*?)\"\,')
        floor = re.compile(r'\"post_no"\:(.*?)\,')
        post_id = re.compile(r'\"post_id\"\:(.*?)\,')
        comment_num = re.compile(r'\"comment_num\"\:(.*?)\,')
        document.add_paragraph('date:' + date.findall(Infmn)[0])
        document.add_paragraph('floor:' + floor.findall(Infmn)[0])
        document.add_paragraph('post_id:' + post_id.findall(Infmn)[0])
        document.add_paragraph('comment_num:' + comment_num.findall(Infmn)[0])
        document.add_paragraph('contents:')
        neirong = (soup.find_all('div', class_='l_post j_l_post l_post_bright  ')[i]).find \
            ('div', class_='d_post_content j_d_post_content  clearfix')
        for child in neirong:
            document.add_paragraph(str(child))
        document.add_paragraph('---------楼层间隔符----------')

if int(num_page[1].text) == 1:
    pass
else:
    for item in range(int(num_page[1].text)):
        # 首先要再载入;
        urls = url + '?pn=' + num_page[1].text
        page = requests.get(urls, header)
        soup1 = BeautifulSoup(page.text, 'lxml')
        num_tiezi = soup1.find_all('div', class_='l_post j_l_post l_post_bright ')
        for i in range(len(num_tiezi)):
            Infmn = soup1.find_all('div', class_='l_post j_l_post l_post_bright ')[i].get('data-field')
            date = re.compile(r'\"date\"\:\"(.*?)\"\,')
            floor = re.compile(r'\"post_no"\:(.*?)\,')
            post_id = re.compile(r'\"post_id\"\:(.*?)\,')
            comment_num = re.compile(r'\"comment_num\"\:(.*?)\,')
            document.add_paragraph('date:' + date.findall(Infmn)[0])
            document.add_paragraph('floor:' + floor.findall(Infmn)[0])
            document.add_paragraph('post_id:' + post_id.findall(Infmn)[0])
            document.add_paragraph('comment_num:' + comment_num.findall(Infmn)[0])
            document.add_paragraph('contents:')
            neirong = (soup1.find_all('div', class_='l_post j_l_post l_post_bright  ')[i]).find \
                ('div', class_='d_post_content j_d_post_content  clearfix')
            for child in neirong:
                document.add_paragraph(str(child))
            document.add_paragraph('---------楼层间隔符----------')

# save file
document.save(soup.title.contents[0] + '.docx')
